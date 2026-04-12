import json
from docx import Document

def add_line(doc, label, value):
    if value:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(" " + str(value))

def generate():
    with open("report.json") as f:
        data = json.load(f)

    doc = Document("master_template.docx")

    # Clear template content (optional)
    doc._body.clear_content()

    # HEADER
    doc.add_heading("Endemic Environmental Services", 0)
    doc.add_heading("Daily Bio-monitoring Report", 1)

    add_line(doc, "Date:", data.get("reportDate"))
    add_line(doc, "Day:", data.get("reportDay"))
    add_line(doc, "Project:", data.get("projectNameNumber"))
    add_line(doc, "Bio-monitor(s):", data.get("biomonitors"))

    # ACTIVITY
    doc.add_heading("Bio-monitoring Activity", 2)
    add_line(doc, "Start Time:", data.get("bioStartTime"))
    add_line(doc, "End Time:", data.get("bioEndTime"))
    add_line(doc, "Notes:", data.get("bioNotes"))

    # WEATHER
    temp = f"{data.get('temperatureStart','')} - {data.get('temperatureEnd','')}"
    wind = f"{data.get('windSpeedStart','')} - {data.get('windSpeedEnd','')}"

    doc.add_heading("Weather Conditions", 2)
    add_line(doc, "Temperature:", temp)
    add_line(doc, "Wind Speed:", wind)
    add_line(doc, "Conditions:", data.get("weatherConditions"))

    # BIOLOGICAL RESOURCES
    doc.add_heading("Biological Resources", 2)

    for entry in data.get("biologicalResources", []):
        add_line(doc, "Species:", entry.get("speciesName"))
        add_line(doc, "GPS:", entry.get("gps"))
        add_line(doc, "Notes:", entry.get("notes"))
        doc.add_paragraph("")

    doc.save("output.docx")
    print("Created output.docx")

if __name__ == "__main__":
    generate()
